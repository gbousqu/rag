# import databutton as db
import re
import io
from io import BytesIO
from typing import Tuple, List
import pickle

from langchain.docstore.document import Document
# from langchain_community.embeddings import OpenAIEmbeddings
from langchain_openai import OpenAIEmbeddings

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from pypdf import PdfReader
from docx import Document as DocxDocument
# import textract

# import faiss
def parse_txt(file: BytesIO, filename: str) -> Tuple[List[str], str]:
    text = file.read().decode('latin-1')
    return [text], filename

def parse_docx(file: BytesIO, filename: str) -> Tuple[List[str], str]:
    doc = DocxDocument(file)
    output = []
    for para in doc.paragraphs:
        output.append(para.text)
    return output, filename


def parse_pdf(file: BytesIO, filename: str) -> Tuple[List[str], str]:
    pdf = PdfReader(file)
    output = []
    for page in pdf.pages:
        text = page.extract_text()
        text = re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)
        text = re.sub(r"(?<!\n\s)\n(?!\s\n)", " ", text.strip())
        text = re.sub(r"\n\s*\n", "\n\n", text)
        output.append(text)
    return output, filename


def text_to_docs(text: List[str], filename: str) -> List[Document]:
    if isinstance(text, str):
        text = [text]
    page_docs = [Document(page_content=page) for page in text]
    for i, doc in enumerate(page_docs):
        doc.metadata["page"] = i + 1

    doc_chunks = []
    for doc in page_docs:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=4000,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            chunk_overlap=0,
        )
        chunks = text_splitter.split_text(doc.page_content)
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk, metadata={"page": doc.metadata["page"], "chunk": i}
            )
            doc.metadata["source"] = f"{doc.metadata['page']}-{doc.metadata['chunk']}"
            doc.metadata["filename"] = filename  # Add filename to metadata
            doc_chunks.append(doc)
    return doc_chunks


def docs_to_index(docs, openai_api_key):
    index = FAISS.from_documents(docs, OpenAIEmbeddings(openai_api_key=openai_api_key))
    return index


# def get_index_for_pdf(pdf_files, pdf_names, openai_api_key):
#     documents = []
#     for pdf_file, pdf_name in zip(pdf_files, pdf_names):
#         text, filename = parse_pdf(BytesIO(pdf_file), pdf_name)
#         documents = documents + text_to_docs(text, filename)
#     index = docs_to_index(documents, openai_api_key)
#     return index


def get_index_for_files(files, file_names, openai_api_key):
    documents = []
    for file, file_name in zip(files, file_names):
        if file_name.endswith('.pdf'):
            text, filename = parse_pdf(BytesIO(file), file_name)
        elif file_name.endswith('.txt'):
            text, filename = parse_txt(BytesIO(file), file_name)
        elif file_name.endswith('.docx'):
            text, filename = parse_docx(BytesIO(file), file_name)
        # elif file.name.endswith('.doc'):
        #     text, filename = parse_doc(file, file_name)
        else:
            raise ValueError(f"Unsupported file type: {file.name}")

        documents = documents + text_to_docs(text, filename)

    index = docs_to_index(documents, openai_api_key)
    return index